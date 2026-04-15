import re
import io
from typing import Optional

import pandas as pd
from docx import Document


def _split_emails(raw: str) -> list[str]:
    """Split a cell that may contain multiple emails separated by , or ;"""
    if not raw:
        return []
    parts = re.split(r"[;,]", str(raw))
    return [p.strip() for p in parts if p.strip() and "@" in p]


def extract_from_excel(file_bytes: bytes) -> list[dict]:
    """
    Parse Key Participants table from whiteboard Excel.

    The table has a merged header row containing 'Key Participants'
    followed by rows: Column A = team name, Column B = email(s).
    """
    df = pd.read_excel(io.BytesIO(file_bytes), header=None)

    # Find the row index of the "Key Participants" header
    start_row = None
    for idx, row in df.iterrows():
        row_str = " ".join(str(v) for v in row.values if pd.notna(v)).lower()
        if "key participants" in row_str:
            start_row = idx + 1  # data starts on the next row
            break

    if start_row is None:
        return []

    participants = []
    for idx in range(start_row, len(df)):
        row = df.iloc[idx]
        team = str(row.iloc[0]).strip() if pd.notna(row.iloc[0]) else ""
        email_raw = str(row.iloc[1]).strip() if len(row) > 1 and pd.notna(row.iloc[1]) else ""

        # Stop if both columns are empty (end of table)
        if not team and not email_raw:
            break
        # Skip rows that look like section headers (no @ in email column)
        if not email_raw or "@" not in email_raw:
            continue

        emails = _split_emails(email_raw)
        if emails:
            participants.append({"team": team, "emails": emails})

    return participants


def extract_from_docx(file_bytes: bytes) -> list[dict]:
    """
    Parse Attendees table from existing MIR Word document.

    The attendees table has columns: email | Functional Area
    """
    doc = Document(io.BytesIO(file_bytes))
    attendees = []

    for table in doc.tables:
        # Look for a table whose first header cell contains "Attendees"
        if not table.rows:
            continue
        header_text = table.rows[0].cells[0].text.strip().lower()
        if "attendees" not in header_text:
            continue
        # Data rows start at index 1
        for row in table.rows[1:]:
            if len(row.cells) < 2:
                continue
            email = row.cells[0].text.strip()
            area = row.cells[1].text.strip()
            if "@" in email:
                attendees.append({"email": email, "functional_area": area})

    return attendees


def extract_mir_metadata(file_bytes: bytes) -> dict:
    """
    Extract INC number, PRB number, title, description, resolution,
    business impact from an existing MIR Word document.
    Returns a dict with whatever fields could be found.
    """
    doc = Document(io.BytesIO(file_bytes))
    data = {}

    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            key = cells[0].text.strip().lower()
            value = cells[1].text.strip()

            if "title" in key and "title" not in data:
                data["title"] = value
            elif "incident number" in key and "inc_number" not in data:
                data["inc_number"] = value
            elif "problem number" in key and "prb_number" not in data:
                data["prb_number"] = value

    # Extract bold paragraph fields (Description, Resolution, Business Impact)
    for para in doc.paragraphs:
        text = para.text.strip()
        for field, key in [
            ("Description:", "description"),
            ("Resolution:", "resolution"),
            ("Business Impact:", "business_impact"),
            ("Summary:", "summary"),
        ]:
            if text.startswith(field) and key not in data:
                data[key] = text[len(field):].strip()

    return data

import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _bold_para(cell, text: str, size: int = 10):
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)


def _add_header_row(table, text: str, col_count: int, bg: str = "4B0082"):
    row = table.add_row()
    cell = row.cells[0]
    # Merge across all columns
    for i in range(1, col_count):
        cell = cell.merge(row.cells[i])
    cell.text = text
    _set_cell_bg(cell, bg)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)


def build_mir_docx(mir_data: dict) -> bytes:
    """
    Build a clean MIR Word document from the generated data.
    Mirrors the 3-page FedEx MIR structure.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── Title block ───────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run("Global IT Service Management")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    heading = doc.add_heading("Major Incident Review", 0)
    heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x1F, 0x6B)

    doc.add_paragraph()

    # ── Info table ────────────────────────────────────────────────────────────
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = "Table Grid"

    def add_info_row(label, value):
        row = info_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value or ""
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.size = Pt(10)

    add_info_row("Title", mir_data.get("title", ""))
    add_info_row("Incident Number", mir_data.get("inc_number", ""))
    add_info_row("Problem Number", mir_data.get("prb_number", ""))

    doc.add_paragraph()

    # ── Incident detail paragraphs ────────────────────────────────────────────
    for label, key in [
        ("Incident number", "inc_number"),
        ("Configuration Item", "ci"),
        ("Incident opened", "inc_opened"),
        ("Incident resolved", "inc_resolved"),
        ("Duration degraded", "duration_degraded"),
    ]:
        if mir_data.get(key):
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            run.font.size = Pt(10)
            p.add_run(mir_data[key]).font.size = Pt(10)

    doc.add_paragraph()

    for label, key in [
        ("Description", "description"),
        ("Resolution", "resolution"),
        ("Business Impact", "business_impact"),
    ]:
        if mir_data.get(key):
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            run.font.size = Pt(10)
            p.add_run(mir_data[key]).font.size = Pt(10)

    doc.add_paragraph()

    # ── Impact Heading ────────────────────────────────────────────────────────
    if mir_data.get("impact_heading"):
        p = doc.add_paragraph()
        run = p.add_run("Impact Heading: ")
        run.bold = True
        run.font.size = Pt(10)
        vh = p.add_run(mir_data["impact_heading"])
        vh.bold = True
        vh.font.size = Pt(10)
        vh.font.color.rgb = RGBColor(0x1F, 0x1F, 0x6B)

    doc.add_paragraph()

    # ── Attendees table ───────────────────────────────────────────────────────
    attendees = mir_data.get("attendees", [])
    if attendees:
        att_table = doc.add_table(rows=1, cols=2)
        att_table.style = "Table Grid"
        hdr = att_table.rows[0].cells
        _bold_para(hdr[0], "Attendees")
        _bold_para(hdr[1], "Functional Area")
        _set_cell_bg(hdr[0], "D3D3D3")
        _set_cell_bg(hdr[1], "D3D3D3")

        for entry in attendees:
            for email in entry.get("emails", []):
                row = att_table.add_row()
                row.cells[0].text = email
                row.cells[1].text = entry.get("team", "")
                for c in row.cells:
                    for run in c.paragraphs[0].runs:
                        run.font.size = Pt(9)

        doc.add_paragraph()

    # ── Questions and Observations ────────────────────────────────────────────
    questions = mir_data.get("questions", [])
    if questions:
        qo_table = doc.add_table(rows=1, cols=1)
        qo_table.style = "Table Grid"
        hdr_cell = qo_table.rows[0].cells[0]
        _bold_para(hdr_cell, "Questions and Observations")
        _set_cell_bg(hdr_cell, "D3D3D3")

        for i, q in enumerate(questions, 1):
            row = qo_table.add_row()
            row.cells[0].text = f"{i}. {q}"
            for run in row.cells[0].paragraphs[0].runs:
                run.font.size = Pt(10)

        doc.add_paragraph()

    # ── Action table ──────────────────────────────────────────────────────────
    act_table = doc.add_table(rows=4, cols=2)
    act_table.style = "Table Grid"
    _bold_para(act_table.rows[0].cells[0], "Action to be taken")
    _bold_para(act_table.rows[0].cells[1], "Action owner with target date")
    _set_cell_bg(act_table.rows[0].cells[0], "D3D3D3")
    _set_cell_bg(act_table.rows[0].cells[1], "D3D3D3")

    doc.add_paragraph()

    # ── Page break ────────────────────────────────────────────────────────────
    doc.add_page_break()

    # ── Page 3: CAPA ──────────────────────────────────────────────────────────
    opp_table = doc.add_table(rows=4, cols=2)
    opp_table.style = "Table Grid"
    _bold_para(opp_table.rows[0].cells[0], "Opportunities for improvement")
    _bold_para(opp_table.rows[0].cells[1], "Action owner with target date")
    _set_cell_bg(opp_table.rows[0].cells[0], "D3D3D3")
    _set_cell_bg(opp_table.rows[0].cells[1], "D3D3D3")

    doc.add_paragraph()

    ptask_table = doc.add_table(rows=2, cols=2)
    ptask_table.style = "Table Grid"
    _bold_para(ptask_table.rows[0].cells[0], "PTASKS")
    _bold_para(ptask_table.rows[0].cells[1], "PTASK owner")
    _set_cell_bg(ptask_table.rows[0].cells[0], "D3D3D3")
    _set_cell_bg(ptask_table.rows[0].cells[1], "D3D3D3")

    doc.add_paragraph()

    rc_table = doc.add_table(rows=2, cols=1)
    rc_table.style = "Table Grid"
    _bold_para(rc_table.rows[0].cells[0], "Root Cause Summary")
    _set_cell_bg(rc_table.rows[0].cells[0], "D3D3D3")
    rc_table.rows[1].cells[0].text = mir_data.get("root_cause_summary", "")

    doc.add_paragraph()

    capa_table = doc.add_table(rows=2, cols=1)
    capa_table.style = "Table Grid"
    _bold_para(capa_table.rows[0].cells[0], "Corrective and Preventive Actions")
    _set_cell_bg(capa_table.rows[0].cells[0], "D3D3D3")
    capa_table.rows[1].cells[0].text = mir_data.get("capa", "")

    doc.add_paragraph()

    # ── Footer note ───────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SECURITY CLASSIFICATION: Confidential")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)

    # ── Serialize to bytes ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
